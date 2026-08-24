import io
import json
import logging
from typing import Dict, Any, List
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from backend.app.database.connection import get_db

logger = logging.getLogger(__name__)

def set_rtl(paragraph):
    """Sets Word paragraph direction to Right-to-Left (RTL)."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

class DocumentExporter:
    """Exports translated documents preserving page numbering, headings, and RTL typography."""

    @classmethod
    def export_project_to_docx(cls, project_id: str) -> io.BytesIO:
        doc = Document()
        
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT name FROM projects WHERE id = ?;", (project_id,))
            p_row = cursor.fetchone()
            proj_name = p_row["name"] if p_row else "Tarjuman Export"

            # Title
            title_p = doc.add_paragraph()
            set_rtl(title_p)
            r = title_p.add_run(f"ترجمہ: {proj_name}")
            r.bold = True
            r.font.size = Pt(18)

            # Get documents
            cursor.execute("SELECT * FROM documents WHERE project_id = ? ORDER BY filename ASC;", (project_id,))
            docs = cursor.fetchall()

            for d in docs:
                h_p = doc.add_paragraph()
                set_rtl(h_p)
                hrun = h_p.add_run(f"\n--- کتاب / دستاویز: {d['filename']} ---")
                hrun.bold = True

                cursor.execute("""
                SELECT page_number, chunk_index, source_text, COALESCE(final_urdu, target_urdu, '') as urdu, primary_model
                FROM chunks 
                WHERE document_id = ? 
                ORDER BY page_number ASC, chunk_index ASC;
                """, (d["id"],))
                chunks = cursor.fetchall()

                current_page = 0
                for c in chunks:
                    if c["page_number"] != current_page:
                        current_page = c["page_number"]
                        p_mark = doc.add_paragraph()
                        set_rtl(p_mark)
                        prun = p_mark.add_run(f"« صفحہ {current_page} »")
                        prun.italic = True

                    urdu_text = c["urdu"].strip()
                    if urdu_text:
                        p = doc.add_paragraph()
                        set_rtl(p)
                        p.add_run(urdu_text)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    @classmethod
    def export_project_to_json(cls, project_id: str) -> str:
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM projects WHERE id = ?;", (project_id,))
            proj = dict(cursor.fetchone() or {})

            cursor.execute("SELECT * FROM documents WHERE project_id = ?;", (project_id,))
            docs = [dict(d) for d in cursor.fetchall()]

            for d in docs:
                cursor.execute("SELECT * FROM chunks WHERE document_id = ? ORDER BY page_number, chunk_index;", (d["id"],))
                d["chunks"] = [dict(c) for c in cursor.fetchall()]

            proj["documents"] = docs
            return json.dumps(proj, indent=2, ensure_ascii=False)

    @classmethod
    def export_project_to_txt(cls, project_id: str, bilingual: bool = False) -> str:
        lines = []
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM documents WHERE project_id = ? ORDER BY filename ASC;", (project_id,))
            docs = cursor.fetchall()
            for d in docs:
                lines.append(f"==================================================")
                lines.append(f"DOCUMENT: {d['filename']}")
                lines.append(f"==================================================\n")

                cursor.execute("""
                SELECT page_number, chunk_index, source_text, COALESCE(final_urdu, target_urdu, '') as urdu, primary_model, qa_status
                FROM chunks 
                WHERE document_id = ? 
                ORDER BY page_number ASC, chunk_index ASC;
                """, (d["id"],))
                chunks = cursor.fetchall()
                for c in chunks:
                    lines.append(f"[Page {c['page_number']} - Chunk {c['chunk_index']}] (Model: {c['primary_model']} | QA: {c['qa_status']})")
                    if bilingual:
                        lines.append(f"ARABIC:\n{c['source_text']}\n")
                        lines.append(f"URDU:\n{c['urdu']}\n")
                    else:
                        lines.append(f"{c['urdu']}\n")
                    lines.append("-" * 40)
        return "\n".join(lines)
